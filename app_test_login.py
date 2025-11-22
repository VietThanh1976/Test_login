import streamlit as st
import speech_recognition as sr
import io
import os
# Thư viện cho file DOCX
from docx import Document 
# Thư viện xử lý audio file
import librosa 
import soundfile as sf 
# Thư viện ghi âm từ micro trình duyệt
from streamlit_mic_recorder import mic_recorder 
# Thư viện đăng nhập chuyên nghiệp
import streamlit_authenticator as stauth
# Thư viện đọc file YAML
import yaml
from yaml.loader import SafeLoader

def transcribe_audio_from_file_path(file_path):
    """Sử dụng SpeechRecognition để chuyển đổi file WAV thành văn bản."""
    r = sr.Recognizer()
    try:
        with sr.AudioFile(file_path) as source:
            audio = r.record(source) 
        text = r.recognize_google(audio, language="vi-VN")
        return text
    # ... (các khối except giữ nguyên)
    except sr.UnknownValueError:
        return "Không thể nhận dạng giọng nói từ tệp âm thanh này."
    except sr.RequestError as e:
        return f"Lỗi kết nối hoặc API: {e}"
    except Exception as e:
        return f"Lỗi xử lý tệp: {e}"

def process_uploaded_file(uploaded_file):
    """Xử lý và chuyển đổi file đã tải lên."""
    st.session_state.last_transcription_text = ""
    temp_input_path = "temp_input_audio" + os.path.splitext(uploaded_file.name)[1]
    temp_wav_path = "temp_converted_audio.wav"
    
    # ... (Logic xử lý file dùng librosa và soundfile giữ nguyên)
    try:
        with open(temp_input_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        y, sr_librosa = librosa.load(temp_input_path, sr=None) 
        sf.write(temp_wav_path, y, sr_librosa)
        
        st.info("Đang nhận dạng giọng nói...")
        result_text = transcribe_audio_from_file_path(temp_wav_path)
        
        st.session_state.last_transcription_text = result_text

    except Exception as e:
        st.session_state.last_transcription_text = f"Lỗi xử lý tệp: {e}"
    finally:
        if os.path.exists(temp_input_path):
            os.remove(temp_input_path)
        if os.path.exists(temp_wav_path):
            os.remove(temp_wav_path)

def create_docx(text, filename="transcribed_document.docx"):
    """Tạo một file DOCX từ văn bản đã chuyển đổi."""
    document = Document()
    document.add_heading('Văn bản đã chuyển đổi', 0)
    document.add_paragraph(text)

    docx_io = io.BytesIO()
    document.save(docx_io)
    docx_io.seek(0)
    return docx_io.read(), filename

# 1. ĐỌC FILE CẤU HÌNH YAML
try:
    with open('config.yaml') as file:
        config = yaml.load(file, Loader=SafeLoader)
except FileNotFoundError:
    st.error("Lỗi: Không tìm thấy file 'config.yaml'.")
    st.stop()

# 2. KHỞI TẠO AUTHENTICATOR (Truyền toàn bộ dictionary config)
authenticator = stauth.Authenticate(
    config['credentials'],
    config['cookie']['name'],
    config['cookie']['key'],
    config['cookie']['expiry_days']
)
# 3. Khởi tạo Session State cho ứng dụng chính ---
r = sr.Recognizer()
if 'audio_buffer' not in st.session_state:
    st.session_state.audio_buffer = None
if 'last_transcription_text' not in st.session_state:
    st.session_state.last_transcription_text = ""
if 'last_audio_data' not in st.session_state:
    st.session_state.last_audio_data = None
    
# Gọi hàm login. Đặt form ở cột chính ('main')
name, authentication_status, username = authenticator.login('Đăng nhập', 'main')

# 3. Xử lý logic theo trạng thái
if authentication_status:
    # --- Đã Đăng nhập thành công ---
    # Hiển thị nút Đăng xuất ở thanh bên (sidebar)
    with st.sidebar:
        st.success(f"Chào mừng, {name}!")
        authenticator.logout('Đăng xuất', 'main') 
        
    # --- NỘI DUNG ỨNG DỤNG CHÍNH ĐẶT Ở ĐÂY ---
    st.title("🎤 Ứng Dụng Chuyển Giọng Nói Thành Văn Bản")
    # ... (Phần code xử lý chuyển đổi giọng nói của bạn) ...
    method = st.radio(
        "Chọn phương thức nhập liệu:",
        ('Tải lên File Âm thanh', 'Ghi âm trực tiếp từ Micro')
    )

    ### PHƯƠNG THỨC 1: Tải lên File Âm thanh
    if method == 'Tải lên File Âm thanh':
        uploaded_file = st.file_uploader(
            "Tải lên tệp âm thanh (.wav, .mp3, etc.):",
            type=['wav', 'mp3', 'ogg', 'flac']
        )
        if uploaded_file is not None:
            if st.button('🚀 Chuyển đổi File thành Văn bản'):
                with st.spinner('Đang tải và xử lý file...'):
                    process_uploaded_file(uploaded_file)
                
    ### PHƯƠNG THỨC 2: Ghi âm trực tiếp từ Micro
    elif method == 'Ghi âm trực tiếp từ Micro':
        st.subheader("🎙️ Ghi Âm Trực Tiếp")
        st.caption("Ghi âm bằng micro của trình duyệt.")

        audio_data = mic_recorder(
            start_prompt="Bắt đầu Ghi Âm",
            stop_prompt="Dừng Ghi Âm",
            key='mic_recorder',
            format="wav"
        )

        if audio_data:
            st.session_state.audio_buffer = audio_data['bytes']
            st.session_state.last_audio_data = audio_data['bytes']
            st.audio(st.session_state.audio_buffer, format='audio/wav') 
            
            st.download_button(
                label="⬇️ Tải xuống File Âm thanh (.wav)",
                data=st.session_state.last_audio_data,
                file_name="ghi_am_mic.wav",
                mime="audio/wav"
            )
        
        if st.session_state.audio_buffer is not None:
            if st.button('✅ Chuyển đổi Giọng nói'):
                st.session_state.last_transcription_text = ""
                
                temp_wav_path = "mic_recording_temp.wav"
                
                try:
                    with open(temp_wav_path, "wb") as f:
                        f.write(st.session_state.audio_buffer)

                    with st.spinner('Đang nhận dạng giọng nói...'):
                        result_text = transcribe_audio_from_file_path(temp_wav_path)
                    
                    st.session_state.last_transcription_text = result_text

                except Exception as e:
                    st.session_state.last_transcription_text = f"Lỗi xử lý: {e}"
                finally:
                    if os.path.exists(temp_wav_path):
                        os.remove(temp_wav_path)

    # --- Hiển thị Kết quả và Tùy chọn Tải xuống (Chung) ---
    if st.session_state.last_transcription_text:
        st.markdown("---")
        st.subheader("✅ Văn bản đã chuyển đổi:")
        
        st.text_area("Kết quả:", st.session_state.last_transcription_text, height=250)

        if "Không thể" not in st.session_state.last_transcription_text and "Lỗi" not in st.session_state.last_transcription_text:
            
            col1, col2 = st.columns(2)
            
            docx_bytes, docx_filename = create_docx(st.session_state.last_transcription_text)
            col1.download_button(
                label="💾 Tải xuống MS Word (.docx)",
                data=docx_bytes,
                file_name=docx_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            col2.download_button(
                label="📝 Tải xuống Văn bản thuần (.txt)",
                data=st.session_state.last_transcription_text.encode('utf-8'),
                file_name="transcribed_text.txt",
                mime="text/plain"
            )

elif authentication_status == False:
    # --- Đăng nhập thất bại ---
    st.error('Tên đăng nhập/Mật khẩu không chính xác')

elif authentication_status == None:
    # --- Chưa đăng nhập (Lần đầu truy cập) ---
    st.title("Vui lòng Đăng nhập để sử dụng Ứng dụng")
    st.info('Nhập tên người dùng và mật khẩu của bạn để tiếp tục.')



